# utils/screenshot_manager.py

import os
from datetime import datetime
from docx import Document
from docx.shared import Inches

class ScreenshotManager:
    def __init__(self, test_name: str, app_name: str):
        self.test_name = test_name
        self.app_name = app_name
        self.screenshots = []  # (description, filepath)

        timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
        current_date = datetime.now().strftime("%Y-%m-%d")

        # Prepare paths
        self.doc_filename = f"{test_name}_{timestamp}.docx"
        self.output_folder = os.path.join("OutPut", app_name, current_date)
        os.makedirs(self.output_folder, exist_ok=True)

    def add_screenshot(self, driver, description: str):
        timestamp = datetime.now().strftime("%H-%M-%S")
        safe_description = "".join(c if c.isalnum() or c in (' ', '_') else '_' for c in description)
        filename = f"{safe_description}_{timestamp}.png"
        filepath = os.path.join(self.output_folder, filename)
        try:
            driver.save_screenshot(filepath)
            self.screenshots.append((description, filepath))
        except Exception as e:
            print(f" Could not take screenshot: {e}")

    def generate_doc(self):
        if not self.screenshots:
            print(" No screenshots to add to Word document.")
            return

        try:
            doc = Document()
            doc.add_heading(f'Test Report: {self.test_name}', 0)

            for desc, path in self.screenshots:
                doc.add_paragraph(desc)
                try:
                    doc.add_picture(path, width=Inches(6))
                except Exception as e:
                    doc.add_paragraph(f"(Screenshot not available or error: {e})")

                doc.add_paragraph("\n")

            doc_path = os.path.join(self.output_folder, self.doc_filename)
            doc.save(doc_path)
            print(f" Word document created: {doc_path}")
        except Exception as e:
            print(f" Failed to generate Word document: {e}")
